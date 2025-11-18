"""
Parser de archivos DOCX usando python-docx.
Con soporte opcional para OCR usando docling.
"""

import io
from typing import Any, Dict, Union


def parse_docx(
    file_content: Union[bytes, io.BytesIO], use_ocr: bool = False
) -> Union[str, Dict[str, Any]]:
    """
    Extrae texto de un archivo DOCX.

    Si use_ocr=True, usa docling con OCR para mejor detección de imágenes.

    Args:
        file_content: Contenido del archivo DOCX en bytes o BytesIO
        use_ocr: Si True, usa docling con OCR para detectar fotos

    Returns:
        Si use_ocr=False: Texto extraído del DOCX (str)
        Si use_ocr=True: Dict con 'text', 'has_photo', 'images_count', 'metadata'

    Raises:
        ValueError: Si no se puede parsear el DOCX
    """
    # Si OCR está habilitado, usar docling
    if use_ocr:
        try:
            from parsing.ocr import parse_with_docling

            return parse_with_docling(
                file_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except ImportError:
            print("⚠️ Docling no disponible. Cayendo a método tradicional sin OCR.")
            print("   Para usar OCR, instala: pip install docling")
            # Continuar con método tradicional
        except Exception as e:
            print(f"⚠️ Error con OCR, usando método tradicional: {str(e)}")
            # Continuar con método tradicional

    # Método tradicional
    try:
        from docx import Document

        # Asegurar que tenemos BytesIO
        if isinstance(file_content, bytes):
            file_stream = io.BytesIO(file_content)
        else:
            file_stream = file_content

        # Abrir documento
        doc = Document(file_stream)

        text_parts = []

        # Extraer párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Preservar bullet points
                style_name = paragraph.style.name.lower() if paragraph.style else ""

                # Detectar si es item de lista
                is_list_item = any(
                    [
                        "list" in style_name,
                        paragraph.text.strip().startswith(("•", "-", "*", "·", "○")),
                    ]
                )

                if is_list_item and not paragraph.text.strip().startswith(
                    ("•", "-", "*")
                ):
                    text_parts.append("• " + paragraph.text.strip())
                else:
                    text_parts.append(paragraph.text.strip())

        # Extraer tablas
        for table in doc.tables:
            table_text = _extract_table_text(table)
            if table_text.strip():
                text_parts.append(f"\n[Tabla]\n{table_text}\n")

        if not text_parts:
            raise ValueError("El documento DOCX no contiene texto extraíble")

        text = "\n\n".join(text_parts)

        # Si use_ocr está activado, retornar en formato dict
        if use_ocr:
            return {
                "text": text,
                "has_photo": False,  # Método tradicional no detecta fotos
                "images_count": 0,
                "metadata": {},
            }

        return text

    except Exception as e:
        raise ValueError(f"Error parseando DOCX: {str(e)}")


def _extract_table_text(table) -> str:
    """
    Extrae texto de una tabla de documento Word.

    Args:
        table: Objeto tabla de python-docx

    Returns:
        Texto de la tabla formateado
    """
    rows_text = []

    for row in table.rows:
        cells_text = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                cells_text.append(cell_text)

        if cells_text:
            rows_text.append(" | ".join(cells_text))

    return "\n".join(rows_text)


def get_docx_metadata(file_content: Union[bytes, io.BytesIO]) -> dict:
    """
    Extrae metadatos del documento DOCX.

    Args:
        file_content: Contenido del archivo DOCX

    Returns:
        Diccionario con metadatos
    """
    try:
        from docx import Document

        if isinstance(file_content, bytes):
            file_stream = io.BytesIO(file_content)
        else:
            file_stream = file_content

        doc = Document(file_stream)
        core_props = doc.core_properties

        return {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

    except Exception:
        return {}
